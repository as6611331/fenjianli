import requests, time, os, random,re,csv,xlwt,json,pymysql,datetime
from docx import Document
from bs4 import BeautifulSoup
from selenium import webdriver
from configparser import ConfigParser
from concurrent.futures import ThreadPoolExecutor,ProcessPoolExecutor #线程池，进程池
from shutil import copyfile

class downloader(object):
    def __init__(self):

        # 符合html条件对文件位置
        self.url_all=[]
        self.cookie = ''
        self.data_D_max = 100
        self.data_D_min = 0
        self.data_U_max = 300
        self.data_U_min = 0
        self.condition={}
        self.account=''
        self.account_password=''

        # 数据库启动
        self.sql_status=True
        self.host='192.168.0.41'
        self.user='root'
        self.password=''
        self.db='it-data'
        self.port=3306
        self.charset='utf8mb4'

        # 纷简历-3 doc
        self.url_fenjianli_3 = []
        self.url_fenjianli_3_title=['简历更新时间','简历编号','姓名','性别','手机号码','年龄','电子邮箱','学历','婚姻状况','工作年限','现居住地','户籍','期望行业','期望职业','期望地点','期望薪资','工作性质','目前状态','自我评价','工作经历','项目经历','教育经历','语言能力','培训经历','专业技能','证书','简历来源','创建时间']
        self.url_fenjianli_3_datas=[]

        # 纷简历-4 html 最详细
        self.url_fenjianli_4 = []
        self.url_fenjianli_4_title=['更新时间','简历编号','姓名','性别','手机号码','年龄','电子邮件','教育程度','工作年限','婚姻状况','职业状态','国籍','所在地','户籍','期望行业','期望职位','期望地点','期望薪资','工作经历','项目经历','教育经历','培训经历','专业技能','语言能力','自我评价','所获证书','简历来源','创建时间']
        self.url_fenjianli_4_datas=[]

        # 新猎场-1
        self.url_xinliechang_1_title = ['姓名','手机号','期望职业','性别','年龄','期望薪资','工作年限','学历','现居住地','期望地点','电子邮箱','目前状态','自我评价','工作经历','创建时间']
        self.url_xinliechang_1_datas = []

        # 新猎场-2
        self.url_xinliechang_2_title = ["简历名称","姓名","性别","出生年月","身高","籍贯","婚否","工作经验","学历","职位性质","意向行业","意向工作地区","意向薪资","教育经历","工作经历","语言能力","证书","手机","邮箱","居住地","意向职位","自我描述","添加时间","渠道来源","专业技能","培训经历","目前状态","项目经历",]
        self.url_xinliechang_2_datas = []

        # 汇总数据
        self.days = []
        self.numbers = 0
        self.statistics_title = {'日期':0,'北京': 110000, '深圳': 440300, '上海': 310000, '广州': 440100, '成都': 510100, '杭州': 330100, '重庆': 500000,'天津': 120000}
        # self.statistics_title = {'日期':0,'天津': 120000}
        self.statistics_data = []

        self.conversion_situation={'正确简历':0,'错误简历':0,'其他类型':0}
        self.upload_situation={'上传成功':0,'存在简历':0,'上传失败':0}
        self.download_situation = {'下载成功': 0,'下载失败': 0}

    # 检查数据库重复
    def mysql_judge(self,table_name,operating,data):
        db = pymysql.connect(host=dl.host, user=dl.user, password=dl.password, db=dl.db, port=dl.port,charset=dl.charset)  # 连接数据库
        cur=db.cursor() # 获取一个游标

        if ('resume_id' in data.keys()):
            resume_id = data['resume_id']
        else:
            resume_id = data['简历编号']
            phone_number = data['手机号码']

        if table_name=='fenjianli_id':
            cur.execute("select `resume_id` from {0} where `resume_id`='{1}'".format('fenjianli_id', resume_id))
            ret = cur.fetchone()

            if operating=='select':
                return ret
            elif operating=='insert':
                if ret == None:
                    downdate = time.strftime('%Y-%m-%d', time.localtime(time.time()))
                    upload = [resume_id, downdate, dl.account, str(json.dumps(dl.condition, ensure_ascii=False))]
                    format = str(',%s' * len(upload))[1:]
                    cur.execute('insert into {0} values({1})'.format('fenjianli_id', format), upload)  # 插入数据

        elif table_name=='fenjianli_html' or table_name=='fenjianli_doc':
            cur.execute("select `手机号码` from {0} where `手机号码`='{1}'".format(table_name, phone_number))
            ret = cur.fetchone()
            if phone_number in str(ret):
                sql="DELETE FROM {0} where `手机号码`='{1}'".format(table_name,phone_number) # 删除数据
                cur.execute(sql)
                format = str(',%s' * len(data))[1:]
                cur.execute('insert into {0} values({1})'.format(table_name, format), list(data.values()))  # 插入数据
            else:
                format = str(',%s' * len(data))[1:]
                cur.execute('insert into {0} values({1})'.format(table_name, format), list(data.values()))  # 插入数据

        # 提交
        db.commit()
        # 关闭指针对象
        cur.close()
        # 关闭连接对象
        db.close()

    # 获得cookie
    def get_cookies2(self,status=''):

        cookie_status=os.path.isfile(".\cookie.txt")

        if cookie_status==False or status=='登录失效':
            login = 'http://www.fenjianli.com/login'
            diver = webdriver.Chrome()
            diver.get(login)
            diver.find_element_by_id('linkPwd').click()
            diver.find_element_by_name('partner').send_keys(dl.account)
            diver.find_element_by_name('pwd').send_keys(dl.account_password)
            if dl.account_password != '':
                time.sleep(0.5)
                diver.find_element_by_css_selector("button[data-type=\"pwd-login\"]").click()

            while True:
                time.sleep(1)
                try:
                    dl.cookie = diver.get_cookies()[1]['value']
                    print(dl.cookie)
                    break
                except:
                    pass
            diver.quit()
            with open('cookie.txt', 'w', encoding='UTF-8') as f:
                f.write(dl.cookie)
        else:
            with open('cookie.txt', 'r', encoding='UTF-8') as f:
                dl.cookie = f.read()

    # 读取下载剩余简历数
    def get_score(self):
        url = 'http://www.fenjianli.com/user'
        headers = {
            'Accept': 'application/json, text/javascript, */*; q=0.01',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Length': '0',
            'Host': 'www.fenjianli.com',
            'Origin': 'http://www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/share',
            'X-Requested-With': 'XMLHttpRequest',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36',
        }
        cookies = {'fid': dl.cookie}
        html = requests.post(url, cookies=cookies, headers=headers)
        html = json.loads(html.text)
        htmls = str(html['data']['usable_download_time'])
        return htmls

    # 自动新建文件夹
    def makedirs(self,path):
        if not os.path.exists(path):
            os.makedirs(path)

    # 读取配置
    def get_conf(self,operating=''):
        cp = ConfigParser()
        cp.read(".\config.ini",encoding='utf-8-sig')

        # 读取数据库信息
        dl.sql_status = True if cp.get("mysql_db", "sql_status")=='True' else False
        dl.host = cp.get("mysql_db", "host")
        dl.port = int(cp.getint("mysql_db", "port"))
        dl.db = cp.get("mysql_db", "db")
        dl.user = cp.get("mysql_db", "user")
        dl.password = cp.get("mysql_db", "password")
        dl.charset = cp.get("mysql_db", "charset")

        # 下载条件数据
        if operating=='下载':
            dicts={}
            dl.account = cp.get("search_condition", "account")
            dl.account_password = cp.get("search_condition", "account_password")
            dl.data_D_max = int(cp.get("search_condition", "data_D_max"))
            for i in ['keywords','city','age','degree','sex','salarys','update','hideDownloaded','page']:
                data=cp.get("search_condition", i)
                if data != '':
                    dicts.update({i: data})
            dl.condition = dicts

            print('下载账号：{0} 下载数量：{1} 下载条件：{2}'.format(dl.account,dl.data_D_max,dl.condition))

        elif operating=='上传':
            dl.account = cp.get("search_condition", "account")
            dl.account_password = cp.get("search_condition", "account_password")
            dl.data_U_max = int(cp.get("search_condition", "data_U_max"))

    '''--------------------转换程序--------------------'''

    # csv模块导出csv
    def csv_to_csv(self, name,title, datas):
        headers = list(title)
        with open(name + '.csv', 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.DictWriter(f, headers, delimiter=',', quotechar='"')
            writer.writeheader()
            # print(datas)
            for row in datas:
                writer.writerow(row)

    # xlwt模块导出xls
    def xlwt_to_xls(self,name,title,datas):
        #65535行
        workbook = xlwt.Workbook(encoding='UTF-8')
        worksheet = workbook.add_sheet(name)

        for i in range(len(title)):
            worksheet.write(0, i, label=title[i])

        for n in range(1,len(datas)+1):
            for i in range(len(title)):
                worksheet.write(n,i,str(list(datas[n-1].values())[i]))

        workbook.save(name+'.xls')

    # 纷简历-3清洗程序 doc
    def get_url_fenjianli_3(self,url):
        try:
            # print(url)
            dicts = dict.fromkeys(dl.url_fenjianli_3_title, "")

            doc_tables=Document(url).tables

            dicts['简历更新时间']=doc_tables[0].cell(0, 0).text.replace('简历更新时间:', '')
            dicts['简历编号']=doc_tables[0].cell(0, 1).text.replace('简历编号：', '')

            dicts['姓名']=doc_tables[2].cell(1, 1).text
            dicts['性别']=doc_tables[2].cell(1, 3).text
            dicts['手机号码']=doc_tables[2].cell(2, 1).text
            dicts['年龄']=doc_tables[2].cell(2, 3).text
            dicts['电子邮箱']=doc_tables[2].cell(3, 1).text
            dicts['学历']=doc_tables[2].cell(3, 3).text
            dicts['婚姻状况'] = doc_tables[2].cell(4, 1).text
            dicts['工作年限'] = doc_tables[2].cell(4, 3).text
            dicts['现居住地'] = doc_tables[2].cell(5, 1).text
            dicts['户籍'] = doc_tables[2].cell(5, 3).text

            dicts['自我评价'] = "".join([paragraph.text for paragraph in Document(url).paragraphs if paragraph.text != ''])

            professional_skills_lists = []
            for i in range(len(doc_tables)):
                doc_tables_text=doc_tables[i].cell(0, 0).text
                # print(doc_tables_text)

                if '职业发展意向' in doc_tables_text:
                    # print("职业发展意向：")
                    doc_tables_data=doc_tables[i + 1]
                    for rows in range(len(doc_tables_data.rows)):
                        if '期望行业'in doc_tables_data.cell(rows, 0).text:
                            dicts['期望行业'] = str(doc_tables_data.cell(rows, 1).text.split(';')).replace("'",'"')
                        elif '期望职业'in doc_tables_data.cell(rows, 0).text:
                            dicts['期望职业'] = str(doc_tables_data.cell(rows, 1).text.split(';')).replace("'",'"')
                        elif '期望地点'in doc_tables_data.cell(rows, 0).text:
                            dicts['期望地点'] = str(doc_tables_data.cell(rows, 1).text.split('-')).replace("'",'"')
                        elif '期望薪资'in doc_tables_data.cell(rows, 0).text:
                            dicts['期望薪资'] = doc_tables_data.cell(rows, 1).text
                        elif '工作性质'in doc_tables_data.cell(rows, 0).text:
                            dicts['工作性质'] = doc_tables_data.cell(rows, 1).text
                        elif '目前状态'in doc_tables_data.cell(rows, 0).text:
                            dicts['目前状态'] = doc_tables_data.cell(rows, 1).text

                elif '工作经历' in doc_tables_text:
                    # print("工作经历：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    for rowss in range(len(doc_tables_data.rows)):
                        if len(doc_tables_data.cell(rowss, 0).text) > 6:
                            tables=['任职时间','公司','行业','职位','工作描述']
                            tables_dicts = dict.fromkeys(tables, '')
                            tables_dicts['任职时间']=doc_tables_data.cell(rowss, 0).text
                            tables_dicts['公司'] = doc_tables_data.cell(rowss, 1).text.replace(re.search(re.compile(r'\((\d*(年|个月))+\)'), doc_tables_data.cell(rowss, 1).text.replace('()','(0年)')).group(), '')
                            for rows in range(len(doc_tables_data.rows)):
                                if '行业' in doc_tables_data.cell(rows, 0).text:
                                    tables_dicts['行业'] = doc_tables_data.cell(rows, 1).text
                                elif '职位' in doc_tables_data.cell(rows, 0).text:
                                    tables_dicts['职位'] = doc_tables_data.cell(rows, 1).text
                                elif '工作描述' in doc_tables_data.cell(rows, 0).text:
                                    tables_dicts['工作描述'] = doc_tables_data.cell(rows, 1).text
                            lists.append(tables_dicts)
                            # print(tables_dicts)
                    dicts['工作经历']=str(json.dumps(lists,ensure_ascii=False))

                elif '项目经历' in doc_tables_text:
                    # print("项目经历：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    for rowss in range(len(doc_tables_data.rows)):
                        if len(doc_tables_data.cell(rowss, 0).text) > 6:
                            tables=['项目时间','项目名称','项目职责','项目描述']
                            tables_dicts = dict.fromkeys(tables, '')
                            tables_dicts['项目时间'] = doc_tables_data.cell(rowss, 0).text
                            tables_dicts['项目名称'] = doc_tables_data.cell(rowss, 1).text.replace(re.search(re.compile(r'\((\d*(年|个月))+\)'), doc_tables_data.cell(rowss, 1).text.replace('()','(0年)')).group(), '').replace(' ','')
                            for rows in range(len(doc_tables_data.rows)):
                                if "项目职责" in doc_tables_data.cell(rows, 0).text:
                                    tables_dicts['项目职责'] = doc_tables_data.cell(rows, 1).text
                                if "项目描述"in doc_tables_data.cell(rows, 0).text:
                                    tables_dicts['项目描述'] = doc_tables_data.cell(rows, 1).text
                            lists.append(tables_dicts)
                            # print(tables_dicts)
                    dicts['项目经历']=str(json.dumps(lists,ensure_ascii=False))

                elif '教育经历' in doc_tables_text:
                    # print("教育经历：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    lens = []
                    for row in doc_tables_data.rows:
                        for cell in row.cells:
                            lens.append(cell.text)

                    for i in range(len(lens)):
                        try:
                            if lens[i][8] =='-':
                                tables = ['就读时间', '学校', '学历', '专业']
                                tables_dicts = dict.fromkeys(tables, '')
                                tables_dicts['就读时间'] = lens[i]
                                r=i+1
                                for n in range(3):
                                    try:
                                        if lens[r] == '学校：':
                                            tables_dicts['学校'] = lens[r + 1]
                                        elif lens[r] == '学历：':
                                            tables_dicts['学历'] = lens[r + 1]
                                        elif lens[r] == '专业：':
                                            tables_dicts['专业'] = lens[r + 1]
                                    except:
                                        pass
                                    finally:
                                        r=r+2
                                # print(tables_dicts)
                                lists.append(tables_dicts)
                        except:
                            pass
                    # print(lists)
                    dicts['教育经历'] = str(json.dumps(lists,ensure_ascii=False))

                elif '语言能力' in doc_tables_text:
                    # print("语言能力：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    for rows in range(len(doc_tables_data.rows))[1:]:
                        tables=['语言','读写情况','听说情况']
                        tables_dicts = dict.fromkeys(tables, '')
                        split = doc_tables_data.cell(rows, 0).text.split(':')
                        tables_dicts['语言'] = split[0]
                        tables_dicts['读写情况'] = split[1]
                        tables_dicts['听说情况'] = split[2]
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['语言能力'] = str(json.dumps(lists,ensure_ascii=False))

                elif '培训经历' in doc_tables_text:
                    # print("培训经历：")
                    doc_tables_data=doc_tables[i + 1]
                    for rows in range(len(doc_tables_data.rows)):
                        try:
                            if doc_tables_data.cell(rows, -1).text[11] == '-':
                                tables=['培训时间','培训机构','培训课程']
                                tables_dicts = dict.fromkeys(tables, '')
                                tables_dicts['培训时间'] = doc_tables_data.cell(rows, -1).text.replace('-','/').replace(' /',' -')
                                tables_dicts['培训机构'] = doc_tables_data.cell(rows, 0).text
                                tables_dicts['培训课程'] = doc_tables_data.cell(rows, 1).text
                                professional_skills_lists.append(tables_dicts)
                                # print(tables_dicts)
                        except:
                            pass

                elif '专业技能' in doc_tables_text:
                    # print("专业技能：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    for rows in range(len(doc_tables_data.rows))[1:]:
                        lists.append(doc_tables_data.cell(rows, 0).text.replace('\\xa0',' ').replace('\\u3000',' ').replace('"','').replace("'",''))
                    # print(lists)
                    dicts['专业技能'] = str(lists).replace("'",'"')

                elif '证书' in doc_tables_text:
                    # print("证书：")
                    doc_tables_data=doc_tables[i + 1]
                    lists = []
                    for rows in range(len(doc_tables_data.rows))[1:]:
                        tables=['获得时间','证书名称']
                        tables_dicts = dict.fromkeys(tables, '')
                        tables_dicts['获得时间'] = doc_tables_data.cell(rows, 0).text
                        tables_dicts['证书名称'] = doc_tables_data.cell(rows, 1).text
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['证书'] = str(json.dumps(lists,ensure_ascii=False))

            if len(professional_skills_lists)>0:
                dicts['培训经历'] = str(json.dumps(professional_skills_lists,ensure_ascii=False))

            dicts['简历来源'] ='纷简历doc'
            dicts['创建时间'] = time.strftime('%Y-%m-%d',time.localtime(time.time()))

            # print(dicts)
            dl.url_fenjianli_3_datas.append(dicts)
        except:
            dl.conversion_situation['错误简历'] += 1
            print('错误简历-纷简历-3：', os.path.basename(url))
            copyfile(url, '.\data-转换\错误简历\%s'%(os.path.basename(url)))
        else:
            dl.conversion_situation['正确简历'] += 1
            print('正确简历-纷简历-3：', os.path.basename(url))
            try:
                dl.get_url_xinliechang_1(dicts)
            except:
                print('错误转换-新猎场-1：',os.path.basename(url))
            try:
                dl.get_url_xinliechang_2(dicts)
            except:
                print('错误转换-新猎场-2：',os.path.basename(url))
            if dl.sql_status == True:
                dl.mysql_judge('fenjianli_id', 'insert', dicts)
                dl.mysql_judge('fenjianli_doc', '', dicts)

    # 纷简历-4清洗程序 html
    def get_url_fenjianli_4(self, url):
        try:
            with open(url, 'r', encoding='UTF-8') as f:
                soup = BeautifulSoup(f, 'lxml')
            dicts = dict.fromkeys(dl.url_fenjianli_4_title, '')
            try:
                dicts['简历编号'] = soup.find('input').get('value')
            except:
                pass
            soup = soup.find(class_='menu-box')

            pointer = soup.findAll(class_='update')[0].text.replace('更新时间：', '')
            dicts['更新时间'] = pointer

            pointer = soup.find(class_='cont relative')
            label = pointer.findAll('label')
            col = pointer.findAll(class_='col')
            for i in range(len(label)):
                if '姓名' in label[i].text:
                    dicts['姓名'] = col[i].text
                elif '性别' in label[i].text:
                    dicts['性别'] = col[i].text
                elif '手机号码' in label[i].text:
                    dicts['手机号码'] = col[i].text
                elif '年龄' in label[i].text:
                    dicts['年龄'] = col[i].text
                elif '电子邮件' in label[i].text:
                    dicts['电子邮件'] = col[i].text
                elif '教育程度' in label[i].text:
                    dicts['教育程度'] = col[i].text
                elif '工作年限' in label[i].text:
                    dicts['工作年限'] = col[i].text
                elif '婚姻状况' in label[i].text:
                    dicts['婚姻状况'] = col[i].text
                elif '职业状态' in label[i].text:
                    dicts['职业状态'] = col[i].text
                elif '国籍' in label[i].text:
                    dicts['国籍'] = col[i].text
                elif '所在地' in label[i].text:
                    dicts['所在地'] = col[i].text
                elif '户籍' in label[i].text:
                    dicts['户籍'] = col[i].text
                elif '期望行业' in label[i].text:
                    dicts['期望行业'] = str(col[i].text.split(';')).replace("'",'"')
                elif '期望职位' in label[i].text:
                    dicts['期望职位'] = str(col[i].text.replace('\n', '').replace('\\xa0\\xa0',';').replace('\\xa0',' ').replace('\\u3000',' ').split(';')).replace("'",'"')
                elif '期望地点' in label[i].text:
                    dicts['期望地点'] = str(col[i].text.split('-')).replace("'",'"')
                elif '期望薪资' in label[i].text:
                    dicts['期望薪资'] = col[i].text.replace(' ', '').replace('\n', '')

            float = soup.select('section[class="board"] span[class="float-left"]')
            cout=0
            for n in float:
                #0
                if '工作经历' in n.text:
                    pointer = soup.find(id='workexp_anchor').select('div[class="exp"] > table > tbody > tr')
                    tables = ['任职时间', '公司', '公司性质', '公司规模', '公司行业', '职位', '所在部门', '职责']
                    lists = []
                    for i in pointer:
                        tables_dicts = dict.fromkeys(tables, '')
                        tables_dicts['任职时间'] = i.find(class_='times').text

                        tables_dicts['公司'] = i.findAll(class_='table table-noborder table-form')[0].find(
                            class_="section-content").text.replace(' ', '').replace('\n', '')
                        try:
                            get_dey = re.search(re.compile(r'（(\d*(年|个月))+）'), tables_dicts['公司']).group()
                        except:
                            get_dey = ''
                        tables_dicts['公司'] = tables_dicts['公司'].replace(get_dey, '')

                        comp_info = i.findAll(class_='table table-noborder table-form')[0].findAll(class_='comp-info')
                        for td in comp_info:
                            if '公司性质：' in td.text:
                                tables_dicts['公司性质'] = td.text.replace('\n', '').replace(' ', '').replace('公司性质：', '')
                            elif '公司规模：' in td.text:
                                tables_dicts['公司规模'] = td.text.replace('\n', '').replace(' ', '').replace('公司规模：', '')
                            elif '公司行业：' in td.text:
                                tables_dicts['公司行业'] = td.text.replace('\n', '').replace(' ', '').replace('公司行业：', '')

                        tables_dicts['职位'] = i.findAll(class_='table table-noborder table-form')[1].find(
                            class_="section-content").text.replace(' ', '').replace('\n', '')
                        comp_info = i.findAll(class_='table table-noborder table-form')[1].findAll('tr')
                        for td in comp_info:
                            if '所在部门：' in td.text:
                                tables_dicts['所在部门'] = td.text.replace('\n', '').replace(' ', '').replace('所在部门：', '')
                            elif '职责：' in td.text:
                                tables_dicts['职责'] = td.text.replace('\n', '').replace(' ', '').replace('职责：', '')

                        lists.append(tables_dicts)
                    dicts['工作经历'] = str(json.dumps(lists, ensure_ascii=False))
                #1
                elif '项目经历' in n.text:
                    # print(1)
                    tables = ['项目时间', '项目名称', '项目简介', '项目职责']
                    lists = []
                    pointer = soup.select('section[class="board"] div[class="exp"]')[1]
                    pointer=[i for i in pointer.find("tbody").children if i != '\n']
                    for i in pointer:
                        tables_dicts = dict.fromkeys(tables, '')
                        tables_dicts['项目时间']=i.find(class_='times').text
                        tables_dicts['项目名称'] =i.select('table thead tr th')[0].text
                        k=0
                        for s in i.select('table tbody tr th'):
                            if '项目简介：' in s.text:
                                tables_dicts['项目简介']=i.select('table tbody tr td')[k].text.replace('\n', '').replace(' ', '')
                                k+=1
                            elif '项目职责：'in s.text:
                                tables_dicts['项目职责'] = i.select('table tbody tr td')[k].text.replace('\n', '').replace(' ', '')
                                k += 1
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['项目经历'] = str(json.dumps(lists,ensure_ascii=False))
                    # print(pointer)
                #2
                elif '教育经历' in n.text:
                    # print(2)
                    tables = ['就读时间', '学校', '学历', '专业']
                    lists = []
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    pointer = pointer.select('tbody tr')
                    for i in pointer:
                        pointers =i.select('td')
                        # print(pointers)
                        tables_dicts = dict.fromkeys(tables, '')
                        for s in pointers:
                            if '专业：' in s.text:
                                tables_dicts['专业']=s.text.replace('专业：', '')
                            elif '学历：' in s.text:
                                tables_dicts['学历'] = s.text.replace('学历：', '')
                            else:
                                tables_dicts['就读时间'] =re.search(re.compile(r'\d+\/+\d+( - )+(\d+\/+\d*|至今)'), s.text).group()
                                tables_dicts['学校'] =s.text.replace(tables_dicts['就读时间'],'')
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['教育经历'] = str(json.dumps(lists,ensure_ascii=False))
                    # print(pointer)
                    cout+=1
                #3
                elif '培训经历' in n.text:
                    # print(3)
                    tables = ['培训时间', '培训机构', '培训课程']
                    lists = []
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    pointer = pointer.select('tbody tr')
                    for i in pointer:
                        tables_dicts = dict.fromkeys(tables, '')
                        pointers=i.select('td')
                        lens=len(pointers)
                        if lens==1:
                            tables_dicts['培训时间'] = pointers[0].text.replace('-','/').replace(' /',' -')
                        else:
                            tables_dicts['培训时间'] = pointers[0].text.replace('-','/').replace(' /',' -')
                            tables_dicts['培训机构'] = pointers[1].text
                            tables_dicts['培训课程'] = pointers[2].text
                        lists.append(tables_dicts)
                        # print(tables_dicts)

                    dicts['培训经历'] = str(json.dumps(lists,ensure_ascii=False))
                    # print(lists)
                    # print(pointer)
                    cout+=1
                #4
                elif '专业技能' in n.text:
                    # print(4)
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    dicts['专业技能'] = pointer.text.replace(' ', '').replace('\\xa0',' ').replace('\\u3000',' ').replace("'",'').replace('"','').split('\n')
                    dicts['专业技能'] = str([i for i in dicts['专业技能'] if i != '']).replace("'",'"')
                    # print(dicts['专业技能'] )
                    # print(pointer)
                    cout+=1
                #5
                elif '语言能力' in n.text:
                    # print(5)
                    tables = ['语言', '读写情况', '听说情况']
                    lists = []
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    pointer=pointer.findAll(class_="language")
                    for i in pointer:
                        tables_dicts = dict.fromkeys(tables, '')
                        i=i.text.replace(' ', '').replace('：', '/').replace('\n', '').split('/')
                        for n in i:
                            if '读写能力'in n:
                                tables_dicts['读写情况'] = n.replace('读写能力', '')
                            elif'听说能力'in n:
                                tables_dicts['听说情况'] = n.replace('听说能力', '')
                            else:
                                tables_dicts['语言']=n
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['语言能力'] = str(json.dumps(lists,ensure_ascii=False))
                    # print(pointer)
                    cout+=1
                #6
                elif '自我评价' in n.text:
                    # print(6)
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    dicts['自我评价'] = pointer.text.replace('\n', '')
                    # print(dicts['自我评价'])
                    # print(pointer)
                    cout+=1
                #7
                elif '所获证书' in n.text:
                    # print(7)
                    tables = ['获得时间', '证书名称']
                    lists = []
                    pointer = soup.select('section[class="board"] div[class="cont"]')[cout]
                    pointer=pointer.findAll('p')
                    for i in pointer:
                        tables_dicts = dict.fromkeys(tables, '')
                        i=i.text.replace('\n', '').replace(' ', '').split('\xa0')
                        i= [n for n in i if n != '']
                        if len(i)==1:
                            tables_dicts['获得时间'] = i[0]
                        else:
                            tables_dicts['获得时间'] = i[0]
                            tables_dicts['证书名称'] = i[1]
                        lists.append(tables_dicts)
                        # print(tables_dicts)
                    dicts['所获证书'] = str(json.dumps(lists,ensure_ascii=False))
                    # print(pointer)
                    cout += 1

            dicts['简历来源'] ='纷简历html'
            dicts['创建时间'] = time.strftime('%Y-%m-%d',time.localtime(time.time()))

            # print(dicts)
            dl.url_fenjianli_4_datas.append(dicts)
        except:
            dl.conversion_situation['错误简历'] += 1
            print('错误简历-纷简历-4：',os.path.basename(url))
            copyfile(url, '.\data-转换\错误简历\%s'%(os.path.basename(url)))
        else:
            dl.conversion_situation['正确简历'] += 1
            print('正确简历-纷简历-4：', os.path.basename(url))
            try:
                dl.get_url_xinliechang_1(dicts)
            except:
                print('错误转换-新猎场-1：',os.path.basename(url))
            try:
                dl.get_url_xinliechang_2(dicts)
            except:
                print('错误转换-新猎场-2：',os.path.basename(url))
            if dl.sql_status == True:
                dl.mysql_judge('fenjianli_id', 'insert', dicts)
                dl.mysql_judge('fenjianli_html', '', dicts)

    # 新猎场-1导入转换
    def get_url_xinliechang_1(self, data):
        dicts = dict.fromkeys(dl.url_xinliechang_1_title, '')
        dicts['姓名']= data['姓名']
        dicts['手机号'] = data['手机号码']
        dicts['性别'] = data['性别']
        dicts['年龄'] = data['年龄'].replace('岁','')
        dicts['期望薪资'] = data['期望薪资']
        dicts['工作年限'] = data['工作年限']
        try:
            dicts['期望地点'] =  ','.join(json.loads(data['期望地点']))
        except:
            pass
        dicts['自我评价'] = data['自我评价']
        try:
            # html
            try:
                dicts['期望职业'] = ','.join(json.loads(data['期望职位']))
            except:
                pass
            dicts['学历'] = data['教育程度']
            dicts['现居住地'] = data['所在地']
            dicts['电子邮箱'] = data['电子邮件']
            dicts['目前状态'] = data['职业状态']
        except:
            # doc
            try:
                dicts['期望职业'] = ','.join(json.loads(data['期望职业']))
            except:
                pass
            dicts['学历'] = data['学历']
            dicts['现居住地'] = data['现居住地']
            dicts['电子邮箱'] = data['电子邮箱']
            dicts['目前状态'] = data['目前状态']

        gzjl=eval(data['工作经历'])
        gzjl_data=''
        n=1
        for i in gzjl:
            gzjl_data=gzjl_data+'{0}.{1}-{2}职位:{3}'.format(n,i['任职时间'].replace('.','/'),i['公司'],i['职位'])
            n+=1
        dicts['工作经历'] = gzjl_data

        dicts['创建时间'] = data['创建时间']
        dl.url_xinliechang_1_datas.append(dicts)
        # print(dicts)

    # 新猎场-2导入转换
    def get_url_xinliechang_2(self, data):
        dicts = dict.fromkeys(dl.url_xinliechang_2_title, '')
        dicts['简历名称']= data['姓名']
        dicts['姓名'] = data['姓名']
        dicts['性别'] = data['性别']
        try:
            dicts['出生年月'] = int(time.strftime("%Y", time.localtime())) - int(data['年龄'].replace('岁', ''))
        except:
            pass
        dicts['籍贯'] = data['户籍']
        dicts['婚否'] = data['婚姻状况']
        dicts['工作经验'] = data['工作年限']

        if 'html' in data['简历来源']:
            dicts['学历'] = data['教育程度']


            if data['所获证书'] != '':
                for i in json.loads(data['所获证书']):
                    dicts['证书'] += '【#】{0}|{1}' .format(i['证书名称'],i['获得时间'])
                dicts['证书'] = dicts['证书'][3:]

            dicts['邮箱'] = data['电子邮件']
            dicts['居住地'] = data['所在地']

            if data['期望职位'] != '':
                dicts['意向职位'] = ','.join(json.loads(data['期望职位']))

            dicts['目前状态'] = data['职业状态']

            if data['工作经历'] != '':
                lists=[]
                for i in json.loads(data['工作经历']):
                    list = dict.fromkeys(["任职时间", "公司", "行业", "职位", "工作描述"], "")
                    list['任职时间'] = i['任职时间']
                    list['公司'] = i['公司']
                    list['行业'] = i['公司行业']
                    list['职位'] = i['职位']
                    list['工作描述'] = i['职责']
                    lists.append(list)

                dicts['工作经历']=str(json.dumps(lists, ensure_ascii=False))

            if data['项目经历'] != '':
                lists=[]
                for i in json.loads(data['项目经历']):
                    list = dict.fromkeys(["项目时间", "项目名称", "项目职责", "项目描述"], "")
                    list['项目时间'] = i['项目时间']
                    list['项目名称'] = i['项目名称']
                    list['项目职责'] = i['项目职责']
                    list['项目描述'] = i['项目简介']
                    lists.append(list)

                dicts['项目经历']=str(json.dumps(lists, ensure_ascii=False))

        else:
            dicts['学历'] = data['学历']
            if data['证书'] != '':
                for i in json.loads(data['证书']):
                    dicts['证书'] += '【#】{0}|{1}' .format(i['证书名称'],i['获得时间'])
                dicts['证书'] = dicts['证书'][3:]

            dicts['邮箱'] = data['电子邮箱']
            dicts['居住地'] = data['现居住地']
            if data['期望职业'] != '':
                dicts['意向职位'] = ','.join(json.loads(data['期望职业']))
            dicts['目前状态'] = data['目前状态']
            dicts['工作经历'] = data['工作经历']
            dicts['项目经历'] = data['项目经历']

        if data['期望行业'] != '':
            dicts['意向行业']='【#】'.join(json.loads(data['期望行业']))

        if data['期望地点'] != '':
            dicts['意向工作地区'] = '/'.join(json.loads(data['期望地点']))

        dicts['意向薪资'] =int(float(data['期望薪资'].split('-')[0].split('K')[0])*1000) if data['期望薪资'] != "面议" else "面议"

        if data['教育经历'] != '':
            for i in json.loads(data['教育经历']):
                dicts['教育经历'] += '【#】{0}|{1}|{2}|{3}'.format(i['就读时间'], i['学校'],i['专业'],i['学历'])
            dicts['教育经历'] = dicts['教育经历'][3:]

        if data['语言能力'] != '':
            for i in json.loads(data['语言能力']):
                if i['语言'] != '':
                    dicts['语言能力'] += '【#】{0}|{1}'.format(i['语言'], i['读写情况'])
            dicts['语言能力'] = dicts['语言能力'][3:]

        if data['专业技能'] != '':
            dicts['专业技能'] = ','.join(json.loads(data['专业技能']))

        dicts['手机'] = data['手机号码']
        dicts['自我描述'] = data['自我评价']
        dicts['添加时间'] = data['创建时间']
        dicts['渠道来源'] = data['简历来源']

        if data['培训经历'] != '':
            lists = []
            for i in json.loads(data['培训经历']):
                list = dict.fromkeys(["培训时间", "培训机构", "培训课程"], "")
                list['培训时间'] = i['培训时间'].replace('-','/').replace(' /',' -')
                list['培训机构'] = i['培训机构']
                list['培训课程'] = i['培训课程']
                lists.append(list)
            dicts['培训经历'] = str(json.dumps(lists, ensure_ascii=False))


        dl.url_xinliechang_2_datas.append(dicts)
        # print(dicts)

    # 简易转换分类控制
    def get_task(self, name, get_url, url, title, datas):

        pool = ThreadPoolExecutor()
        task = [i for i in pool.map(get_url, url)]

        # 利用csv模块导出csv
        # dl.csv_to_csv(name, title, datas)

        # 利用xlwt导出xls
        dl.xlwt_to_xls(name, title, datas)

    # 启动转换程序
    def turn_data_program(self):
        try:
            # 读取配置文件
            if os.path.exists('config.ini') == True:
                dl.get_conf()

            # 启动文件地址提取和转换程序
            for root, dirs, files in os.walk('.\data-转换'):
                for file in files:
                    files=os.path.splitext(file)[1]
                    if files == '.html':
                        dl.url_fenjianli_4.append(os.path.join(root, file))

                    elif files == '.doc':
                        dl.url_fenjianli_3.append(os.path.join(root, file))
                    else:
                        dl.conversion_situation['其他类型'] += 1
                        print('其他类型：', file)

            if len(dl.url_fenjianli_3) == 0 and len(dl.url_fenjianli_4) == 0:
                print('请放入转换文件')
                time.sleep(5)
            else:
                dl.account = '转换'

                if len(dl.url_fenjianli_3) != 0:
                    dl.get_task('纷简历-3', dl.get_url_fenjianli_3, dl.url_fenjianli_3, dl.url_fenjianli_3_title,dl.url_fenjianli_3_datas)

                if len(dl.url_fenjianli_4) != 0:
                    dl.get_task('纷简历-4', dl.get_url_fenjianli_4, dl.url_fenjianli_4, dl.url_fenjianli_4_title,dl.url_fenjianli_4_datas)

                # 新猎场导入用_1
                if len(dl.url_xinliechang_1_datas) != 0:
                    dl.xlwt_to_xls('新猎场-1', dl.url_xinliechang_1_title, dl.url_xinliechang_1_datas)

                # 新猎场导入用_2
                if len(dl.url_xinliechang_2_datas) != 0:
                    dl.xlwt_to_xls('新猎场-2', dl.url_xinliechang_2_title, dl.url_xinliechang_2_datas)

        except BaseException as e:
            print(e)
        finally:
            print()
            print("正确数量：%d | 错误数量：%d | 其他类型：%d" % (dl.conversion_situation['正确简历'], dl.conversion_situation['错误简历'], dl.conversion_situation['其他类型']))
            input("回车结束程序")

    '''--------------------上传程序--------------------'''

    # 上传文件
    def post_files(self,path):

        #给个随机文件名
        name = str(random.randint(10000000, 100000000))+os.path.splitext(path)[-1]
        url = 'http://www.fenjianli.com/share/upload'
        cookies = {'fid': dl.cookie}
        with open(path, 'rb') as f:

            files = {'file': (name, f, 'application/msword', {'Expires': '0'})}
            success = requests.post(url, files=files, cookies=cookies)
            success=str(json.loads(success.text))
            if '上传成功' in success:
                msg = '上传成功'
                dl.upload_situation['上传成功'] += 1
                dl.data_U_min += 1

            elif '已存在相同简历' in success:
                msg = '存在简历'
                dl.upload_situation['存在简历'] += 1

            elif '登录状态已失效' in success:
                msg = '登录状态已失效'

            else:
                msg = '上传失败'
                dl.upload_situation['上传失败'] += 1
        return msg,path

    # 启动上传程序
    def up_data_program(self):
        try:
            for root, dirs, files in os.walk('.\data-上传'):
                for file in files:
                    files=os.path.splitext(file)[1]
                    splitext=['.doc','.docx','.xls','.xlsx','.pdf','.txt','.html']
                    if files in splitext:
                        dl.url_all.append(os.path.join(root, file))

            if len(dl.url_all) == 0:
                print('请放入上传文件')
                time.sleep(5)

            else:
                if os.path.exists('config.ini') == True:
                    dl.get_conf('上传')
                dl.get_cookies2()

                # 多线程
                pool = ThreadPoolExecutor(5)
                for data_report_1,data_report_2 in pool.map(dl.post_files, dl.url_all):
                    if dl.data_U_min < dl.data_U_max:
                        print(data_report_1 + "：剩余" + dl.get_score() + ' 上传成功' + str(dl.data_U_min))
                        if data_report_1 == '登录状态已失效':
                            print('----------《登录失效请重新登录账户》----------')
                            break
                        os.remove(data_report_2)
                    else:
                        break

        except BaseException as e:
            print(e)
        finally:
            print()
            print("上传成功：%d | 存在简历：%d | 上传失败：%d"  % (dl.upload_situation['上传成功'], dl.upload_situation['存在简历'],dl.upload_situation['上传失败']))
            input("回车结束程序")

    '''--------------------下载程序--------------------'''

    # 搜索条件模块
    def search_condition(self):

        dl.account=input('下载账号：')
        dl.data_D_max=int(input('下载数量：'))

        #条件
        dicts = {}
        keywords = input('关键词：')
        if keywords != '':
            dicts.update({'keywords': keywords})

        city = input('城市编号：')
        if city != '':
            dicts.update({'city': city})

        age = input('年龄：')
        if age != '':
            dicts.update({'age': age})

        degree = input('学历编号：')
        if degree != '':
            dicts.update({'degree': degree})

        sex = input('性别：')
        if sex != '':
            dicts.update({'sex': sex})

        salarys = input('期望薪资：')
        if salarys != '':
            dicts.update({'salarys': salarys})

        update = input('更新日期：')
        if update != '':
            dicts.update({'update': update})

        dicts.update({'hideDownloaded': 1})
        dicts.update({'page': 1})
        dl.condition = dicts

    # 获得简历ID
    def get_resume_id(self,page):
        url = 'http://www.fenjianli.com/search/list'
        cookies = {'fid': dl.cookie}
        headers={
            'Accept': 'application/json, text/plain, */*',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Length': '28',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Host': 'www.fenjianli.com',
            'Origin': 'http://www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/search',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
        }
        dl.condition['page']=page
        payload=dl.condition
        html = requests.post(url, data=payload, cookies=cookies,headers=headers).text

        for i in json.loads(html)['data']['data']:
            if dl.data_D_min < dl.data_D_max:
                dl.search_mysql(i['es_id'])
                # print(i['es_id'])
            else:
                break

    # 判断ID是否存在
    def search_mysql(self,resume_id):

        #检查数据库是否存在这个简历ID
        if dl.sql_status == True:
            judge = dl.mysql_judge('fenjianli_id', 'select', {'resume_id': resume_id})
        else:
            judge = None

        if judge==None:

            # 下载联系方式状态:剩余积分不足/您已经下载过了/success
            exchange = dl.exchange(resume_id)


            if exchange == 'success':
                print('未下载简历：剩下',dl.get_score())

                if dl.data_D_min != 0:
                    time.sleep(random.randint(10, 50)/10)

                #下载html文件
                dl.download_html(resume_id)
                judge_result=dl.down_judge(resume_id)

                # 下载doc文件
                dl.download_doc(resume_id)

                #上传数据库
                if judge_result=='成功':

                    # 上传ID到数据库里
                    if dl.sql_status == True:
                        dl.mysql_judge('fenjianli_id', 'insert', {'resume_id': resume_id})

                    # 下载并上传数据
                    # dl.get_url_fenjianli_3(os.path.getsize(os.path.join('data-下载\doc\\', resume_id + '.doc')))
                    # dl.get_url_fenjianli_4(os.path.getsize(os.path.join('data-下载\html\\', resume_id + '.html')))

                    dl.data_D_min+=1
                    dl.download_situation['下载成功'] += 1

                else:
                    print('下载失败简历：剩下', dl.get_score())
                    dl.download_situation['下载失败'] += 1

            elif exchange == '您已经下载过了':
                print('已下载简历：剩下', dl.get_score())

            elif exchange == '剩余积分不足':
                print('账号积分不足，请上传简历补充积分')
                dl.data_D_min = dl.data_D_max
            else:
                print('登入一次72招浏览器后再来下载')
                dl.data_D_min = dl.data_D_max
        else:
            print('已下载简历：剩下',dl.get_score())

    # 下载联系方式
    def exchange(self,resume_id):
        url = 'http://www.fenjianli.com/resume/download'
        cookies = {'fid': dl.cookie}
        headers={
            'Accept': 'application/json, text/plain, */*',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Length': '31',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Host': 'www.fenjianli.com',
            'Origin': 'http://www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/resume/resumeTemplate?resumeId=' + resume_id + '&keywords=',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
        }
        payload={'resumeId':resume_id}
        html = requests.post(url, data=payload, cookies=cookies,headers=headers).text

        htmlf=json.loads(html)['msg']
        return htmlf

    # 下载简历html数据
    def download_html(self,resume_id):
        url = 'http://www.fenjianli.com/resume/resumeTemplate'
        cookies = {'fid': dl.cookie}
        headers={
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Host': 'www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/resume/detail?resumeIds='+resume_id+'&keywords=',
            'Upgrade-Insecure-Requests': '1',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
        }
        payload = {'resumeId': resume_id}
        html = requests.get(url, params=payload, cookies=cookies,headers=headers).text

        with open('data-下载\html\\' + resume_id + '.html', 'w', encoding='UTF-8') as code:
            code.write(html)

    # 下载简历doc数据
    def download_doc(self,resume_id):
        url = 'http://www.fenjianli.com/resume/export'
        cookies = {'fid': dl.cookie}
        headers={
            'Accept': 'application/json, text/plain, */*',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Length': '41',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Host': 'www.fenjianli.com',
            'Origin': 'http://www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/resume/resumeTemplate?resumeId=' + resume_id + '&keywords=',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
        }
        payload={'resumeId':resume_id,'type':'word'}
        html = requests.post(url, data=payload, cookies=cookies,headers=headers)

        with open('data-下载\doc\\' + resume_id + '.doc', 'wb') as code:
            code.write(html.content)

    # 判断文件是否成功下载
    def down_judge(self,resume_id):
        size = os.path.getsize(os.path.join('data-下载\html\\', resume_id + '.html'))

        if size == 10271:
            print('----------《登录失效请重新登录账户》----------')
            dl.get_cookies2('登录失效')
            dl.download_html(resume_id)
            # time.sleep(random.randint(100, 200) / 10)

        cycle_number = 0
        while size == 5298 and cycle_number <= 3:
            time.sleep(30)
            dl.download_html(resume_id)
            size = os.path.getsize(os.path.join('data-下载\html\\', resume_id + '.html'))
            cycle_number += 1

        if size !=5298:
            return '成功'
        else:
            return '失败'

    # 启动下载程序
    def down_data_program(self):

        try:
            # 确认条件
            while True:
                if os.path.exists('config.ini') == True:
                    dl.get_conf('下载')
                else:
                    dl.search_condition()
                confirm = input('是否确定条件（Y/N）')
                if confirm == 'Y' or confirm == 'y':
                    break
                print()

            dl.get_cookies2()

            #获得每页ID
            for i in range(1,135):
                # print(i)
                if dl.data_D_min < dl.data_D_max:
                    try:
                        dl.get_resume_id(i)
                        time.sleep(random.randint(10, 30) /10)
                    except:
                        n=3
                        while n :
                            try:
                                time.sleep(random.randint(10, 30) / 10)
                                dl.get_resume_id(i)
                                break
                            except:
                                n -= 1
                        if n==0:
                            print('----------《登录失效请重新登录账户》----------')
                            dl.get_cookies2('登录失效')
                            dl.get_resume_id(i)
                            time.sleep(random.randint(10, 30)/10)
                else:
                    break
        except BaseException as e:
            print(e)
        finally:
            #下载结果报告
            print("下载成功：%d | 下载失败：%d" % (dl.download_situation['下载成功'], dl.download_situation['下载失败']))
            input("回车结束程序")

    '''--------------------统计程序--------------------'''

    # 获得简历日期
    def get_resume_days(self, page):
        url = 'http://www.fenjianli.com/search/list'
        cookies = {'fid': dl.cookie}
        headers = {
            'Accept': 'application/json, text/plain, */*',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Length': '28',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Host': 'www.fenjianli.com',
            'Origin': 'http://www.fenjianli.com',
            'Referer': 'http://www.fenjianli.com/search',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
        }
        dl.condition['page'] = page
        payload = dl.condition
        html = requests.post(url, data=payload, cookies=cookies, headers=headers).text
        htmls=[]
        for i in json.loads(html)['data']['data']:
            htmls.append(i['last_date_show'])
        dl.numbers += 1
        return htmls

    # 获得数据切点位置
    def get_data_slice(self,page,data,days,city):
        position = ''
        if days in data and (page == 1 or len(list(set(data))) == 2) and position == '':
            for a in range(len(data)):
                if days == data[a]:
                    position = [page, a + 1]
                    break
        if days in data and page != 1 and position == '':
            for a in range(page-15,page):
                if position == '':
                    while True:
                        try:
                            data = dl.get_resume_days(a)
                            break
                        except:
                            dl.get_cookies2('登录失效')
                    print('城市：{0} 页数：{1} 搜索次数：{2}'.format(city,a,dl.numbers),list(set(data)))

                    for b in range(len(data)):
                        if days == data[b]:
                            position = [a, b + 1]
                            break
                else:
                    break
        return position

    # 数据数量统计
    def get_quantity(self,city):
        pointer = ['',''] #昨天，前天
        for a in range(1, 135 + 15, 15):
            if pointer[1] == '':
                while True:
                    try:
                        data = dl.get_resume_days(a)
                        break
                    except:
                        dl.get_cookies2('登录失效')
                print('城市：{0} 页数：{1} 搜索次数：{2}'.format(city, a, dl.numbers),list(set(data)))

                if pointer[0] == '':
                    pointer[0] = dl.get_data_slice(a, data, dl.days[1],city)
                elif pointer[0] != '' and pointer[1] == '':
                    pointer[1] = dl.get_data_slice(a, data, dl.days[2],city)
            else:
                break

        if pointer[0] == '':
            pointer[0] = [1, 1]

        if pointer[1] == '':
            pointer[1] = [135, 1]
            print("把这个括号里的问题截图给我看（城市：{0} 年龄段：{1}）".format(city, dl.condition['age']))


        pointer[0] = (pointer[0][0] - 1) * 30 + pointer[0][1] - 1
        pointer[1] = (pointer[1][0] - 1) * 30 + pointer[1][1] - 1 - pointer[0]
        dl.statistics_data[0][city] += pointer[0]
        dl.statistics_data[1][city] += pointer[1]
        # print(dicts_1)
        # print(dicts_2)

    # 启动统计程序
    def statistics_data_program(self):
        dl.get_cookies2()
        dl.condition = {'city': '', 'age': '18,40', 'page': '1'}
        for d in [0,-1,-2]:
            dl.days.append((datetime.datetime.now() + datetime.timedelta(days=d)).strftime('%Y-%m-%d'))

        dl.statistics_data.append(dict.fromkeys(dl.statistics_title.keys(), 0))
        dl.statistics_data.append(dict.fromkeys(dl.statistics_title.keys(), 0))

        dl.statistics_data[0]['日期'] = dl.days[0]
        dl.statistics_data[1]['日期'] = dl.days[1]

        for k, v in dl.statistics_title.items():
            if k != '日期':
                dl.condition['city'] = v
                dl.condition['age'] = '18,40'
                while True:
                    try:
                        # 检测是否到达前天以下
                        data = dl.get_resume_days(134)
                        break
                    except:
                        dl.get_cookies2('登录失效')

                # 正式版
                if dl.days[1] in data:
                    for s in ['18,23', '24,25', '26,27', '28,29', '30,31', '32,34', '35,37', '38,40']:
                        dl.condition['age'] = s
                        dl.get_quantity(k)
                else:
                    dl.get_quantity(k)

                # 测试版 检测是否到达前天以下
                # if dl.days[1] in data:
                #     for s in ['18,23', '24,25', '26,27', '28,29', '30,31', '32,34', '35,37', '38,40']:
                #         dl.condition['age'] = s
                #         while True:
                #             try:
                #                 data = dl.get_resume_days(134)
                #                 if dl.days[1] in data:
                #                     print(k, list(set(data)), s,'再间隔')
                #                 else:
                #                     print(k,list(set(data)),s)
                #                 break
                #             except:
                #                 dl.get_cookies2('登录失效')

        dl.xlwt_to_xls('纷简历-数量统计', list(dl.statistics_title), dl.statistics_data)

    '''--------------------直取程序--------------------'''

    # 检查数据库重复
    def mysql_judge_1(self):
        db = pymysql.connect(host=dl.host, user=dl.user, password=dl.password, db=dl.db, port=dl.port,charset=dl.charset)  # 连接数据库
        cur=db.cursor() # 获取一个游标
        data=cur.execute("select * from {0}".format('fenjianli_doc'))
        datas=cur.fetchall()


        # 提交
        db.commit()
        # 关闭指针对象
        cur.close()
        # 关闭连接对象
        db.close()

        return datas

    def test_mysql(self):
        data = dl.mysql_judge_1()
        for i in data:
            dicts = dict.fromkeys(dl.url_fenjianli_3_title, '')
            for s in range(len(i)):
                dicts[dl.url_fenjianli_3_title[s]] = i[s]
            dl.get_url_xinliechang_2(dicts)

        # 新猎场导入用_2
        if len(dl.url_xinliechang_2_datas) != 0:
            dl.csv_to_csv('新猎场-3', dl.url_xinliechang_2_title, dl.url_xinliechang_2_datas)

if __name__=='__main__':

    dl = downloader()
    print('请选择模式：')
    print('《转换》输入"1" | 《上传》输入"2" | 《下载》输入"3" | 《统计》输入"4"')
    dl.makedirs('.\data-上传')
    dl.makedirs('.\data-下载\html')
    dl.makedirs('.\data-下载\doc')
    dl.makedirs('.\data-转换')
    dl.makedirs('.\data-转换\错误简历')
    select=input()
    if select=='1':
        dl.turn_data_program()
    elif select=='2':
        dl.up_data_program()
    elif select=='3':
        dl.down_data_program()
    elif select=='4':
        dl.statistics_data_program()





